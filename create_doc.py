#!/usr/bin/env python3
"""Generate a Word document containing the system architecture flowchart."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

heading = doc.add_heading("System Architecture Flowchart", level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Raspberry Pi 5 Smart LED Controller")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.italic = True

doc.add_paragraph()

pic = doc.add_picture(
    "/home/user/Embedded-project/system_architecture_flowchart.png",
    width=Inches(6.5)
)
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

caption = doc.add_paragraph(
    "Figure 1: System architecture showing the dual-SPI bus design. "
    "SPI0 connects the nRF24L01+ RF transceiver for remote command capture. "
    "SPI1 drives the SK6812 RGBW LED string through a bidirectional level shifter. "
    "The Raspberry Pi 5 handles packet capture, hex conversion, button map lookup, "
    "and LED state management."
)
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.runs[0].font.size = Pt(9)
caption.runs[0].font.italic = True

output = "/home/user/Embedded-project/system_architecture_flowchart.docx"
doc.save(output)
print(f"Saved to {output}")
